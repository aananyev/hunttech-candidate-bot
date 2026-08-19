"""
CV Parser — извлечение текста из docx и PDF файлов.
"""
from __future__ import annotations

import logging
import subprocess
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: Path) -> str:
    """Извлечь текст из .docx файла."""
    try:
        # На macOS используем textutil (как в навыке hrm-candidate-creation)
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(file_path)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except Exception as e:
        logger.warning("textutil failed, falling back to python-docx: %s", e)

    # Fallback: python-docx
    try:
        import docx
        doc = docx.Document(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text.strip())
        return "\n".join(texts)
    except Exception as e:
        logger.error("python-docx extraction failed: %s", e)
        raise ValueError(f"Не удалось извлечь текст из DOCX: {e}")


def extract_text_from_pdf(file_path: Path) -> str:
    """Извлечь текст из .pdf файла."""
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    texts.append(text.strip())
        return "\n".join(texts)
    except Exception as e:
        logger.error("pdfplumber extraction failed: %s", e)
        raise ValueError(f"Не удалось извлечь текст из PDF: {e}")


def extract_text_from_file(file_path: Path) -> str:
    """Извлечь текст из файла резюме (docx или pdf)."""
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {suffix}. Поддерживаются: .docx, .pdf")


def sanitize_text_for_cv(text: str) -> str:
    """Очистка текста для сохранения в text_cv (HTML-экранирование + <br>)."""
    # HTML-экранирование
    text = text.replace("&", "&").replace("<", "<").replace(">", ">")
    # Переносы строк -> <br>
    text = text.replace("\n", "<br>")
    return text